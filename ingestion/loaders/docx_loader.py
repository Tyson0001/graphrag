"""
DOCX document loader.
"""

import logging
from pathlib import Path
from typing import Optional

from docx import Document

logger = logging.getLogger(__name__)


class DOCXLoader:
    """Loads content from DOCX files."""

    def load(self, file_path: Path) -> Optional[str]:
        """
        Load text content from a DOCX file.

        Args:
            file_path: Path to the DOCX file

        Returns:
            Extracted text content or None if failed
        """
        try:
            doc = Document(str(file_path))

            text_content = []

            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)

            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text)
                    if row_text:
                        text_content.append(" | ".join(row_text))

            if not text_content:
                logger.warning(f"No text content extracted from DOCX: {file_path}")
                return None

            full_text = "\n\n".join(text_content)
            logger.info(f"Successfully loaded DOCX: {file_path}")
            return full_text

        except Exception as e:
            logger.error(f"Failed to load DOCX {file_path}: {e}")
            return None
